"""
Module de traitement des documents
Supporte: PDF, TXT, DOC, DOCX, MD
"""

import os
from typing import List, Dict
import PyPDF2
from docx import Document
import markdown


class DocumentProcessor:
    """Classe pour extraire le texte des différents types de documents"""
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.txt', '.doc', '.docx', '.md']
    
    def process_file(self, filepath: str) -> Dict[str, any]:
        """
        Traite un fichier et extrait son contenu texte
        
        Args:
            filepath: Chemin vers le fichier
            
        Returns:
            Dict contenant le texte, métadonnées et statut
        """
        if not os.path.exists(filepath):
            return {'success': False, 'error': 'Fichier non trouvé'}
        
        filename = os.path.basename(filepath)
        extension = os.path.splitext(filename)[1].lower()
        
        if extension not in self.supported_extensions:
            return {'success': False, 'error': f'Extension {extension} non supportée'}
        
        try:
            if extension == '.pdf':
                text = self._extract_pdf(filepath)
            elif extension == '.txt':
                text = self._extract_txt(filepath)
            elif extension in ['.doc', '.docx']:
                text = self._extract_docx(filepath)
            elif extension == '.md':
                text = self._extract_markdown(filepath)
            else:
                return {'success': False, 'error': 'Type de fichier non reconnu'}
            
            return {
                'success': True,
                'filename': filename,
                'text': text,
                'char_count': len(text),
                'word_count': len(text.split()),
                'extension': extension
            }
        except Exception as e:
            return {'success': False, 'error': f'Erreur lors du traitement: {str(e)}'}
    
    def _extract_pdf(self, filepath: str) -> str:
        """Extrait le texte d'un PDF"""
        text = ""
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_txt(self, filepath: str) -> str:
        """Extrait le texte d'un fichier TXT"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        # Si aucun encodage ne fonctionne, utiliser errors='ignore'
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    
    def _extract_docx(self, filepath: str) -> str:
        """Extrait le texte d'un fichier DOCX"""
        doc = Document(filepath)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extraire aussi le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return "\n".join(text)
    
    def _extract_markdown(self, filepath: str) -> str:
        """Extrait le texte d'un fichier Markdown"""
        with open(filepath, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        # Convertir markdown en HTML puis extraire le texte
        # Pour simplifier, on garde le markdown brut
        return md_content
    
    def process_directory(self, directory: str) -> List[Dict]:
        """
        Traite tous les fichiers supportés dans un dossier
        
        Args:
            directory: Chemin vers le dossier
            
        Returns:
            Liste des résultats de traitement
        """
        results = []
        
        if not os.path.exists(directory):
            return results
        
        for filename in os.listdir(directory):
            filepath = os.path.join(directory, filename)
            
            if os.path.isfile(filepath):
                result = self.process_file(filepath)
                results.append(result)
        
        return results
